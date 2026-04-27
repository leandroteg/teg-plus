"""Parte 3 — Regras, dados, frontend, n8n, integrações, RLS, KPIs, cenário, glossário."""
import sys
sys.path.insert(0, "/home/user/teg-plus/scripts")
from gerar_doc_solicitacao_compras import (  # noqa
    doc, add_heading, add_paragraph, add_bullet, add_table,
    add_code_block, add_callout, page_break,
)
import gerar_doc_parte2  # noqa: F401  - executa parte 2 também
from docx.shared import Pt


# 12. REGRAS DE NEGÓCIO
add_heading(doc, "12. Regras de Negócio e Alçadas", level=1)

add_heading(doc, "12.1 Categorias e Compradores", level=2)
add_table(
    doc,
    ["Categoria", "Comprador", "Aprovador Alçada 1", "Limite Alçada 1"],
    [
        ["MATERIAIS_OBRA", "Lauany", "Welton", "R$ 2.000"],
        ["EPI_EPC", "Lauany", "Welton", "R$ 2.000"],
        ["FERRAMENTAL", "Lauany", "Welton", "R$ 2.000"],
        ["CENTRO_DIST", "Lauany", "Welton", "R$ 2.000"],
        ["AQUISICOES_ESP", "Lauany", "Claudinor", "R$ 2.000"],
        ["FROTA_EQUIP", "Fernando", "Claudinor", "R$ 2.000"],
        ["SERVICOS", "Fernando", "Claudinor", "R$ 2.000"],
        ["LOCACAO", "Fernando", "Claudinor", "R$ 2.000"],
        ["MOBILIZACAO", "Aline", "Welton", "R$ 2.000"],
        ["ALOJAMENTO", "Aline", "Welton", "R$ 2.000"],
        ["ALIMENTACAO", "Aline", "Welton", "R$ 2.000"],
        ["ESCRITORIO", "Aline", "Welton", "R$ 2.000"],
    ],
)

add_heading(doc, "12.2 Regras de Validação", level=2)
add_paragraph(doc, "Na criação:", bold=True)
for item in [
    "Solicitante obrigatório (auto pelo usuário logado).",
    "Obra obrigatória (lookup em obras ativas).",
    "Categoria obrigatória.",
    "Mínimo 1 item; quantidade > 0; unidade obrigatória.",
    "Se urgência ≠ normal, justificativa de urgência obrigatória.",
    "Valor total calculado > 0.",
]:
    add_bullet(doc, item)

add_paragraph(doc, "Na aprovação:", bold=True)
for item in [
    "Token válido, não expirado, não usado.",
    "RC em status 'em_aprovacao' ou 'em_esclarecimento'.",
    "Alçadas seguidas em ordem (não pula níveis).",
]:
    add_bullet(doc, item)

add_paragraph(doc, "Na cotação:", bold=True)
for item in [
    "Mínimo de fornecedores conforme política da categoria.",
    "Quantidade dos itens é READ-ONLY (vem da RC).",
    "Valor total > 0.",
    "Aumento > 50% sobre o estimado dispara alerta (não bloqueia).",
    "Pelo menos 1 fornecedor marcado como selecionado=true.",
]:
    add_bullet(doc, item)

page_break(doc)

# 13. ESTRUTURA DE DADOS
add_heading(doc, "13. Estrutura de Dados (Supabase / PostgreSQL)", level=1)

add_heading(doc, "13.1 Tabelas Principais", level=2)
add_table(
    doc,
    ["Tabela", "Função", "Migration"],
    [
        ["cmp_requisicoes", "Cabeçalho da RC (solicitante, obra, valor, status, alçada)", "001_schema_compras.sql"],
        ["cmp_requisicao_itens", "Linhas/itens da RC", "001_schema_compras.sql"],
        ["cmp_categorias", "12 categorias com políticas e regras de cotação", "004_schema_cotacoes.sql"],
        ["cmp_compradores", "Lauany / Fernando / Aline e suas categorias", "004_schema_cotacoes.sql"],
        ["cmp_cotacoes", "Cotação consolidada (status, fornecedor vencedor, valor)", "004_schema_cotacoes.sql"],
        ["cmp_cotacao_fornecedores", "Cada proposta individual (com itens_precos JSONB)", "004_schema_cotacoes.sql"],
        ["cmp_pedidos", "Pedido de Compra emitido (PO)", "007_fluxo_real.sql"],
        ["cmp_fornecedores", "Cadastro de fornecedores (sync com Omie)", "—"],
        ["apr_aprovacoes", "Tabela unificada de aprovações (multi-módulo)", "011_schema_financeiro.sql + 042"],
        ["alcadas", "Estrutura de níveis e faixas de valor", "001_schema_compras.sql"],
        ["atividades_log", "Auditoria — quem fez o quê e quando", "001_schema_compras.sql"],
        ["fin_contas_pagar", "CP gerada automaticamente após PO", "011_schema_financeiro.sql"],
    ],
)

add_heading(doc, "13.2 Enums Críticos", level=2)
add_code_block(
    doc,
    """status_requisicao: rascunho, pendente, em_aprovacao, em_esclarecimento,
                  devolvida_solicitante, aprovada, rejeitada,
                  em_cotacao, cotacao_enviada, cotacao_aprovada, cotacao_rejeitada,
                  pedido_emitido, em_entrega, entregue,
                  aguardando_pgto, pago, comprada (legado), cancelada

status_aprovacao: pendente, aprovada, rejeitada, expirada, esclarecimento
status_cotacao:   pendente, em_andamento, concluida, cancelada
urgencia_tipo:    normal, urgente, critica
tipo_aprovacao:   requisicao_compra, cotacao, autorizacao_pagamento,
                  minuta_contratual, aprovacao_transporte""",
)

add_heading(doc, "13.3 Functions / RPCs Importantes", level=2)
add_table(
    doc,
    ["Função", "Propósito", "Migration"],
    [
        ["gerar_numero_requisicao()", "Gera número sequencial RC-YYYYMM-NNNNN", "001_schema_compras.sql"],
        ["determinar_alcada(valor)", "Retorna 1-4 conforme faixa de valor", "001_schema_compras.sql"],
        ["get_alerta_cotacao(rc_id)", "Verifica se cotação atende política mínima", "058_alerta_cotacao_politica.sql"],
        ["get_aprovacoes_pendentes_compras()", "Retorna pendências com RC + cotação em 1 query", "027_rpc_aprovacoes_batch.sql"],
    ],
)

add_heading(doc, "13.4 Views (Dashboard)", level=2)
for item in [
    "vw_dashboard_requisicoes — KPIs por status do mês.",
    "vw_requisicoes_completas — RC + alçada + status aprovação atual.",
    "vw_requisicoes_por_obra — KPIs por canteiro.",
    "vw_kpis_compras — total mês, aguardando, aprovadas, rejeitadas, valor, tempo médio aprovação.",
]:
    add_bullet(doc, item)

page_break(doc)

# 14. FRONTEND
add_heading(doc, "14. Frontend — Páginas, Componentes e Hooks", level=1)

add_heading(doc, "14.1 Páginas Principais", level=2)
add_table(
    doc,
    ["Página", "Linhas", "Função"],
    [
        ["NovaRequisicao.tsx", "1.369", "Wizard de criação/edição de RC (4 passos)"],
        ["ListaRequisicoes.tsx", "1.364", "Pipeline com 9 abas por status"],
        ["RequisicaoDetalhe.tsx", "1.144", "Detalhes + timeline + ações de aprovação"],
        ["FilaCotacoes.tsx", "679", "Fila de cotações pendentes do comprador"],
        ["CotacaoForm.tsx", "1.862", "Coleta de propostas + comparativo"],
        ["AprovAi.tsx", "1.856", "Painel mobile de aprovações por token"],
        ["Pedidos.tsx", "1.806", "Gestão de POs (emissão, recebimento, NF, pagamento)"],
    ],
)

add_heading(doc, "14.2 Hooks Principais", level=2)
add_table(
    doc,
    ["Hook", "Tamanho", "Função"],
    [
        ["useRequisicoes.ts", "27 KB", "CRUD de RC + reenvio após devolução"],
        ["useCotacoes.ts", "15 KB", "CRUD de cotação + alerta de política"],
        ["useAprovacoes.ts", "50 KB", "Pendências, decisões, histórico, KPIs"],
        ["usePedidos.ts", "15 KB", "Emissão, atualização, liberação e baixa"],
        ["useAiParse.ts", "20 KB", "Wrapper Gemini Flash Vision"],
        ["useCategorias.ts", "<1 KB", "Lista 12 categorias com políticas"],
    ],
)

add_heading(doc, "14.3 Componentes Reutilizáveis", level=2)
for item in [
    "FluxoTimeline.tsx — visualiza estados de uma RC do início ao fim.",
    "CotacaoComparativo.tsx — tabela + chart lado a lado de fornecedores.",
    "FornecedorCadastroModal.tsx — cadastro rápido de fornecedor durante cotação.",
    "UploadCotacao.tsx — drag-drop com preview e parse via Gemini Vision.",
    "EmitirPedidoModal.tsx — formulário de emissão de PO.",
    "ItemAutocomplete.tsx — autocomplete dos itens da RC durante cotação.",
    "StatusBadge.tsx — render de status com cores customizadas.",
    "UpperInput.tsx / UpperTextarea.tsx — força UPPERCASE (padrão sistemas legados).",
    "NumericInput.tsx — input com separador de milhar.",
]:
    add_bullet(doc, item)

page_break(doc)

# 15. N8N
add_heading(doc, "15. Automações n8n", level=1)

add_heading(doc, "15.1 Workflows do Módulo", level=2)
add_table(
    doc,
    ["Workflow", "Webhook", "Função"],
    [
        ["TEG+ | Compras - Nova Requisição", "POST /compras/requisicao", "Cria RC + 1ª aprovação + notificação"],
        ["TEG+ | Compras - Processar Aprovação", "POST /compras/aprovacao", "Aprova/rejeita/esclarece + roteia próxima alçada"],
        ["TEG+ | Compras - AI Parse Requisição", "POST /compras/requisicao-ai", "Gemini Flash extrai itens de PDF/foto"],
        ["TEG+ | Compras - Finalizar Cotação", "POST /compras/cotacao/submit", "Salva propostas + cria aprovação financeira"],
        ["TEG+ | Compras - Emitir PO", "POST /compras/pedidos/emit", "Gera PO + CP + notifica fornecedor"],
        ["TEG+ | Omie - Sync Fornecedores", "POST /omie/sync/fornecedores", "Importa cadastro Omie → cmp_fornecedores"],
        ["TEG+ | Omie - Sync Contas a Pagar", "POST /omie/sync/cp", "Bidirecional CP TEG+ ↔ Omie"],
    ],
)

add_heading(doc, "15.2 Notificações por Evento", level=2)
add_table(
    doc,
    ["Evento", "Canal", "Destinatário"],
    [
        ["RC criada", "WhatsApp + e-mail", "Aprovador Alçada 1"],
        ["RC aprovada (parcial)", "WhatsApp", "Próximo aprovador"],
        ["RC totalmente aprovada", "WhatsApp + e-mail", "Comprador da categoria"],
        ["Esclarecimento solicitado", "WhatsApp", "Solicitante"],
        ["RC devolvida", "WhatsApp", "Solicitante"],
        ["Cotação finalizada", "E-mail", "Aprovador / financeiro"],
        ["PO emitida", "E-mail + WhatsApp", "Fornecedor"],
        ["Material recebido", "WhatsApp", "Solicitante + gestor obra"],
        ["NF disponível para pgto", "E-mail", "Tesouraria"],
        ["Pagamento efetivado", "WhatsApp", "Solicitante + gestor"],
    ],
)

page_break(doc)

# 16. INTEGRAÇÕES
add_heading(doc, "16. Integrações Externas", level=1)

add_heading(doc, "16.1 Omie ERP", level=2)
add_paragraph(
    doc,
    "Sincronização bidirecional do módulo financeiro. RC aprovada + PO emitida gera "
    "automaticamente Conta a Pagar (CP) no Omie via endpoint IncluirContaPagar.",
)
add_code_block(
    doc,
    """{
  "integracao": "cmp-pedidos",
  "nCodFornecedor": "<omie-id-fornecedor>",
  "nValor": 16750.00,
  "dData": "2026-04-27",
  "cDescricao": "RC-202604-00123 Siemens Brasil"
}""",
)

add_heading(doc, "16.2 Evolution API (WhatsApp)", level=2)
add_paragraph(
    doc,
    "Self-hosted. Notificações de aprovação chegam via WhatsApp com link token-based — "
    "aprovador clica e decide sem fazer login (token de uso único).",
)

add_heading(doc, "16.3 BrasilAPI (CNPJ)", level=2)
add_paragraph(
    doc,
    "Auto-preenchimento de dados de fornecedor via GET /cnpj/v1/{cnpj} — retorna razão social, "
    "e-mail e telefone para acelerar cadastro durante cotação.",
)

add_heading(doc, "16.4 Gemini Flash Vision (Google AI)", level=2)
add_paragraph(
    doc,
    "Usado em dois pontos: (1) IA Parse na criação de RC (extrai itens de PDF/foto), "
    "(2) Upload Cotação (extrai dados da cotação do fornecedor automaticamente).",
)

page_break(doc)

# 17. SEGURANÇA / RLS
add_heading(doc, "17. Segurança, RLS e Segregação de Funções", level=1)

add_heading(doc, "17.1 Estratégia de Chaves", level=2)
add_table(
    doc,
    ["Chave", "Onde", "Permissão"],
    [
        ["anon_key", "Frontend (Vercel)", "Leitura com RLS restritivo"],
        ["service_role_key", "n8n (backend)", "Escrita sem RLS (confiado)"],
        ["JWT autenticado", "Frontend após login", "Aplicado conforme policies"],
    ],
)

add_heading(doc, "17.2 Políticas RLS (Compras)", level=2)
for item in [
    "cmp_requisicoes — SELECT autenticado; INSERT autenticado; UPDATE/DELETE service_role.",
    "cmp_cotacoes — idem.",
    "apr_aprovacoes — SELECT autenticado (aprovadores precisam ver pendências); INSERT/UPDATE service_role.",
    "cmp_pedidos — SELECT autenticado; UPDATE/DELETE service_role.",
]:
    add_bullet(doc, item)

add_heading(doc, "17.3 RBAC (sys_roles)", level=2)
add_table(
    doc,
    ["Role", "Pode em Compras"],
    [
        ["Administrador", "Tudo (CRUD + aprovar + override)"],
        ["Diretor", "Ver + aprovar (Alçada 3-4)"],
        ["Gestor", "Ver + criar + editar + aprovar (Alçada 1-2)"],
        ["Requisitante", "Criar + ver próprias RCs"],
        ["Visitante", "Apenas leitura"],
    ],
)

add_heading(doc, "17.4 Segregação de Funções (Controles SOX)", level=2)
add_table(
    doc,
    ["Ação", "Pode", "Não Pode"],
    [
        ["Criar RC", "Qualquer usuário autenticado", "—"],
        ["Aprovar RC", "Aprovador da alçada", "O próprio solicitante"],
        ["Coletar cotação", "Comprador da categoria", "Outros compradores"],
        ["Emitir PO", "Comprador ou Gestor", "Solicitante sozinho"],
        ["Receber material", "Almoxarife", "Comprador / Solicitante"],
        ["Liberar pagamento", "Tesouraria", "Comprador / Almoxarife"],
        ["Registrar pagamento", "Tesouraria", "Comprador / Almoxarife"],
    ],
)

page_break(doc)

# 18. KPIs
add_heading(doc, "18. KPIs e Dashboards", level=1)

add_heading(doc, "18.1 Dashboard Compras (vw_kpis_compras)", level=2)
for item in [
    "Total de RCs do mês.",
    "Aguardando aprovação (count e valor).",
    "Aprovadas (count e valor).",
    "Rejeitadas (count).",
    "Valor total movimentado.",
    "Tempo médio de aprovação (em horas).",
]:
    add_bullet(doc, item)

add_heading(doc, "18.2 Dashboard AprovAi", level=2)
for item in [
    "Pendentes por tipo (requisicao_compra / cotacao / autorizacao_pagamento / minuta_contratual / aprovacao_transporte).",
    "Pendentes por alçada (1, 2, 3, 4).",
    "SLA vencido (prazo expirado, destaque vermelho).",
    "Distribuição por janela (hoje / próx. 3 dias / depois).",
]:
    add_bullet(doc, item)

add_heading(doc, "18.3 Fila de Cotações (Comprador)", level=2)
for item in [
    "Cotações vencidas.",
    "Cotações com urgência crítica.",
    "Tempo médio de cotação por comprador.",
    "Taxa de finalização (cotações concluídas / iniciadas).",
]:
    add_bullet(doc, item)

page_break(doc)

# 19. CENÁRIO REAL
add_heading(doc, "19. Cenário Real Comentado", level=1)
add_paragraph(
    doc,
    "Para ilustrar o fluxo end-to-end, considere a seguinte requisição real, simulada com "
    "dados de produção:",
    italic=True,
)
add_callout(
    doc,
    "Cenário",
    "Solicitante: João (técnico de campo)  •  Categoria: MATERIAIS_OBRA  •  Obra: SE Frutal\n"
    "Itens: 500m de cabo XLPE 50mm² + 20 conectores + 1 lote de acessórios\n"
    "Valor estimado: R$ 18.000  •  Urgência: normal  •  Alçada necessária: 2",
)

add_heading(doc, "19.1 Linha do Tempo", level=2)
add_table(
    doc,
    ["Data/Hora", "Evento", "Status RC"],
    [
        ["27/04 10:30", "João cria a RC em NovaRequisicao.tsx; n8n gera RC-202604-00123", "pendente → em_aprovacao"],
        ["27/04 12:00", "Welton (Alçada 1) aprova via AprovAi com link WhatsApp", "em_aprovacao (segue Alçada 2)"],
        ["27/04 13:30", "Laucídio (Alçada 2) aprova com observação sobre cotação", "aprovada → em_cotacao"],
        ["27/04 14:00", "Sistema cria cmp_cotacoes vinculado a Lauany (comprador da categoria)", "em_cotacao"],
        ["29/04 09:00", "Lauany coleta 3 propostas (Siemens, Nexans, Prysmian)", "em_cotacao"],
        ["02/05 16:00", "Lauany finaliza, seleciona Siemens (R$ 16.750, 20 dias, à vista)", "cotacao_aprovada"],
        ["02/05 16:30", "Lauany clica 'Emitir PO'; sistema gera PO-202604-00789 + CP no Omie", "pedido_emitido"],
        ["22/05 14:00", "Almoxarife confere material em SE Frutal e registra NF nº 123456789", "entregue"],
        ["22/05 15:00", "Tesouraria libera + paga via PIX; sistema baixa CP no Omie", "pago ✅"],
    ],
)

add_heading(doc, "19.2 Resultado", level=2)
for item in [
    "Economia de R$ 1.250 vs valor estimado (16.750 vs 18.000).",
    "Prazo cumprido (entrega em 20 dias conforme proposta).",
    "Auditoria completa em atividades_log + apr_aprovacoes.",
    "Ciclo total (criação → pagamento): 26 dias.",
]:
    add_bullet(doc, item)

page_break(doc)

# 20. GLOSSÁRIO
add_heading(doc, "20. Glossário", level=1)
add_table(
    doc,
    ["Termo", "Definição"],
    [
        ["RC", "Requisição de Compras (alias: Solicitação de Compras). Documento que inicia o fluxo."],
        ["PO", "Purchase Order — Pedido de Compra emitido após cotação aprovada."],
        ["CP", "Conta a Pagar — obrigação financeira gerada no Omie quando PO é emitida."],
        ["NF / NF-e", "Nota Fiscal Eletrônica do fornecedor."],
        ["Alçada", "Nível hierárquico de aprovação por valor (1-4)."],
        ["RFQ", "Request For Quotation — solicitação de cotação aos fornecedores."],
        ["AprovAi", "Painel de aprovações multi-módulo (compras, financeiro, contratos, logística)."],
        ["SLA", "Prazo máximo que o aprovador tem para decidir (24h-72h conforme alçada)."],
        ["RBAC", "Role-Based Access Control — permissões por papel."],
        ["RLS", "Row Level Security — segurança em nível de linha do PostgreSQL."],
        ["Token de Aprovação", "Hash único em apr_aprovacoes.token usado para acesso sem login."],
        ["Esclarecimento", "Decisão intermediária: aprovador devolve dúvida sem rejeitar."],
        ["Devolução", "Comprador devolve a RC para o solicitante revisar escopo."],
        ["Split", "Cotação distribuída — itens diferentes para fornecedores diferentes."],
        ["Centro de Custo", "Código contábil que vincula a despesa à obra/área."],
        ["Classe Financeira", "Classificação contábil do gasto (material, serviço, EPI etc.)."],
    ],
)

page_break(doc)

# 21. ANEXOS
add_heading(doc, "21. Anexos", level=1)

add_heading(doc, "21.1 Migrations Supabase Relevantes", level=2)
for item in [
    "001_schema_compras.sql — schema base (cmp_requisicoes, alcadas, atividades_log).",
    "004_schema_cotacoes.sql — cotações + categorias + compradores.",
    "005_public_read_policy.sql — políticas RLS genéricas.",
    "007_fluxo_real.sql — dados reais (categorias, compradores, cmp_pedidos).",
    "011_schema_financeiro.sql — apr_aprovacoes + fin_contas_pagar.",
    "019_esclarecimento_flow.sql — fluxo de esclarecimentos.",
    "025_rls_granular.sql — RBAC granular (fase 2).",
    "027_rpc_aprovacoes_batch.sql — RPCs de pendências.",
    "042_apr_tipo_aprovacao.sql — expansão de tipos.",
    "051_sys_roles_permissoes.sql — estrutura RBAC.",
    "058_alerta_cotacao_politica.sql — RPC de alerta de cotação.",
    "068_rbac_v2_papeis_setores.sql — papéis por setor.",
]:
    add_bullet(doc, item)

add_heading(doc, "21.2 Documentos de Apoio", level=2)
for item in [
    "/home/user/teg-plus/README.md",
    "/home/user/teg-plus/ROADMAP_ERP_WORLD_CLASS.md",
    "/home/user/teg-plus/docs/ARCHITECTURE.md",
    "/home/user/teg-plus/n8n-docs/WORKFLOWS.md",
    "/home/user/teg-plus/docs/plans/2026-03-04-approval-flow-design.md",
]:
    add_bullet(doc, item)

add_heading(doc, "21.3 Convenções de Nomenclatura", level=2)
add_table(
    doc,
    ["Prefixo", "Módulo"],
    [
        ["cmp_", "Compras"],
        ["fin_", "Financeiro"],
        ["apr_", "Aprovações (unificado)"],
        ["est_", "Estoque/Almoxarifado"],
        ["log_", "Logística"],
        ["con_", "Contratos"],
        ["sys_", "Sistema (roles, permissões, configs)"],
        ["vw_", "Views"],
    ],
)

add_paragraph(doc, "")
add_paragraph(
    doc,
    "Documento gerado a partir do código de produção do TEG+ ERP. Para alterações neste fluxo, "
    "atualizar primeiro a documentação correspondente em /home/user/teg-plus/docs/ e depois "
    "os artefatos de código (migrations, frontend, n8n).",
    italic=True,
)

# Salva versão final
final_path = "/home/user/teg-plus/docs/compras/SOLICITACAO_DE_COMPRAS_TEG_PLUS.docx"
doc.save(final_path)
print(f"Documento final gerado em: {final_path}")
