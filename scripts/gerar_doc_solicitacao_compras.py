"""Gera documentação Word editável do fluxo de Solicitação de Compras do TEG+ ERP."""
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- Helpers ----------
def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_heading(doc, text, level=1, color=(0x1F, 0x4E, 0x79)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, header_color="1F4E79", zebra="F2F2F2"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        shade_cell(hdr[i], header_color)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                shade_cell(cells[c_idx], zebra)
            set_cell_borders(cells[c_idx])
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)


def add_callout(doc, title, body, color="DEEBF6"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, color)
    set_cell_borders(cell)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    r2 = p.add_run(body)
    r2.font.size = Pt(10)


def page_break(doc):
    doc.add_page_break()


# ---------- Build doc ----------
doc = Document()

# Margens
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

# Estilo padrão
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# CAPA
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("\n\n\nTEG+ ERP")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Documentação Funcional e Técnica")
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Módulo de Compras — Solicitação de Compras (Requisição de Compras)")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph("\n\n\n\n")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(
    f"Versão 1.0  |  {datetime.now().strftime('%d/%m/%Y')}\n"
    "Escopo: Fluxo end-to-end (Requisição → Aprovação → Cotação → Pedido → Recebimento → Pagamento)\n"
    "Baseado em: codebase de produção `/home/user/teg-plus/`"
)
r.italic = True
r.font.size = Pt(11)

page_break(doc)

# SUMÁRIO
add_heading(doc, "Sumário", level=1)
sumario = [
    "1. Sumário Executivo",
    "2. Objetivo e Escopo",
    "3. Papéis e Responsabilidades",
    "4. Fluxo End-to-End (Visão Macro)",
    "5. Etapa 1 — Criação da Solicitação",
    "6. Etapa 2 — Aprovação Multi-Alçada",
    "7. Etapa 3 — Cotação (RFQ)",
    "8. Etapa 4 — Emissão do Pedido de Compra (PO)",
    "9. Etapa 5 — Recebimento e Conferência",
    "10. Etapa 6 — Pagamento",
    "11. Máquina de Estados (Status)",
    "12. Regras de Negócio e Alçadas",
    "13. Estrutura de Dados (Supabase / PostgreSQL)",
    "14. Frontend — Páginas, Componentes e Hooks",
    "15. Automações n8n",
    "16. Integrações Externas",
    "17. Segurança, RLS e Segregação de Funções",
    "18. KPIs e Dashboards",
    "19. Cenário Real Comentado",
    "20. Glossário",
    "21. Anexos",
]
for item in sumario:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Cm(0.5)

page_break(doc)

# 1. SUMÁRIO EXECUTIVO
add_heading(doc, "1. Sumário Executivo", level=1)
add_paragraph(
    doc,
    "O TEG+ ERP implementa um fluxo de Solicitação de Compras (Requisição de Compras / RC) "
    "completo e multi-nível, abrangendo desde a criação da requisição em campo até o pagamento "
    "ao fornecedor, com automações n8n, validações de alçada por valor e categoria, cotação "
    "obrigatória, AprovAi (painel de aprovações mobile-friendly por token) e integração nativa "
    "com o Omie ERP para o módulo financeiro.",
)

add_paragraph(doc, "Características principais:", bold=True)
for item in [
    "4 alçadas de aprovação por valor (Coordenador → Gerente → Diretor → CEO).",
    "12 categorias de material com políticas customizadas e comprador dedicado.",
    "3 compradores especializados (Lauany, Fernando, Aline) com responsabilidades por categoria.",
    "Cotação obrigatória — 1 a 3 fornecedores conforme valor e regra de categoria.",
    "Fluxo de esclarecimentos — aprovador pode devolver dúvidas ao solicitante sem rejeitar.",
    "AprovAi — painel mobile/tablet com aprovação por token único (link WhatsApp/e-mail, sem login).",
    "IA (Gemini Flash Vision) — extração automática de itens, obra, urgência, categoria a partir de PDF/foto.",
    "Geração automática de Pedido de Compra (PO) e Conta a Pagar (CP) após aprovação final.",
    "Auditoria completa em atividades_log e histórico de aprovações.",
]:
    add_bullet(doc, item)

page_break(doc)

# 2. OBJETIVO E ESCOPO
add_heading(doc, "2. Objetivo e Escopo", level=1)
add_heading(doc, "2.1 Objetivo", level=2)
add_paragraph(
    doc,
    "Padronizar, controlar e auditar o ciclo completo de aquisição de bens e serviços nas obras "
    "e estruturas administrativas da TEG, garantindo segregação de funções, conformidade com "
    "alçadas de aprovação e rastreabilidade contábil/financeira ponta a ponta.",
)

add_heading(doc, "2.2 Escopo", level=2)
for item in [
    "Solicitação de Compras (RC) — criação, edição, cancelamento e versionamento.",
    "Aprovações multi-alçada com fluxo de esclarecimento.",
    "Cotação (RFQ) — coleta e comparativo de propostas de fornecedores.",
    "Emissão e gestão do Pedido de Compra (PO).",
    "Recebimento físico e conferência fiscal (NF-e).",
    "Liberação de pagamento e baixa financeira.",
    "Integração com Omie (Contas a Pagar) e WhatsApp (Evolution API).",
]:
    add_bullet(doc, item)

add_heading(doc, "2.3 Fora do escopo deste documento", level=2)
for item in [
    "Cadastro de obras (módulo EGP — Engenharia/Gestão de Projetos).",
    "Movimentações de estoque puras (módulo Almoxarifado).",
    "Folha de pagamento e mobilização de pessoal (módulo RH).",
    "Conciliação bancária e DRE (módulo Controladoria).",
]:
    add_bullet(doc, item)

page_break(doc)

# 3. PAPÉIS
add_heading(doc, "3. Papéis e Responsabilidades", level=1)
add_paragraph(
    doc,
    "O fluxo envolve diferentes perfis com permissões e ações específicas. A segregação de "
    "funções é controlada por RBAC (sys_roles, sys_role_permissoes) e por validações nos "
    "workflows n8n.",
)

add_table(
    doc,
    ["Papel", "Quem", "Responsabilidade Principal", "Acesso"],
    [
        ["Solicitante", "Técnico de campo / colaborador", "Cria a RC, descreve itens, justifica urgência, responde esclarecimentos", "Cria + edita próprias RCs"],
        ["Comprador", "Lauany / Fernando / Aline", "Recebe RC, coleta cotações, finaliza fornecedor vencedor, emite PO", "Cotações da sua categoria"],
        ["Aprovador Alçada 1", "Welton / Claudinor", "Aprova RCs até R$ 2.000 / valida pertinência técnica", "AprovAi - Alçada 1"],
        ["Aprovador Alçada 2", "Laucídio / Gerente", "Aprova RCs de R$ 2.001 a R$ 25.000", "AprovAi - Alçada 2"],
        ["Aprovador Alçada 3", "Diretor", "Aprova RCs de R$ 25.001 a R$ 100.000", "AprovAi - Alçada 3"],
        ["Aprovador Alçada 4", "CEO", "Aprova RCs acima de R$ 100.000", "AprovAi - Alçada 4"],
        ["Almoxarife", "Equipe almoxarifado", "Confere recebimento físico, registra NF-e, atualiza estoque", "Pedidos.tsx (recebimento)"],
        ["Tesouraria", "Financeiro", "Libera pagamento, registra baixa, concilia com Omie", "Contas a Pagar"],
        ["Admin Compras", "Gerente de Compras", "Override de regras, cancelamentos, visão consolidada", "Acesso total ao módulo"],
    ],
)

page_break(doc)

# 4. FLUXO MACRO
add_heading(doc, "4. Fluxo End-to-End (Visão Macro)", level=1)
add_paragraph(
    doc,
    "O fluxo segue 6 etapas sequenciais, com possibilidade de retorno ao solicitante "
    "(esclarecimento ou devolução) em pontos específicos.",
)

add_code_block(
    doc,
    """
[1] CRIAÇÃO              [2] APROVAÇÃO              [3] COTAÇÃO
   Solicitante     ─►    Alçada 1 → 2 → 3 → 4  ─►  Comprador
   (NovaRequisicao)      (AprovAi)                  (CotacaoForm)
                              │
                              ├─► Esclarecimento ──► volta solicitante
                              └─► Rejeição ────────► encerra

[4] PEDIDO DE COMPRA      [5] RECEBIMENTO           [6] PAGAMENTO
   Comprador        ─►    Almoxarife          ─►   Tesouraria
   (EmitirPedidoModal)    (Pedidos.tsx)            (Contas a Pagar)
                                                         │
                                                         └─► Status final: PAGO
""",
)

add_callout(
    doc,
    "Pontos de retorno do fluxo",
    "• Esclarecimento: aprovador pede esclarecimentos → RC volta a 'em_esclarecimento' → "
    "solicitante responde → volta para mesma alçada.\n"
    "• Devolução: comprador devolve a RC → status 'devolvida_solicitante' → solicitante revisa "
    "escopo → reenvia.\n"
    "• Rejeição: aprovador rejeita → status 'rejeitada' (terminal).",
)

page_break(doc)

print("Parte 1/3 do documento construída...")
doc.save("/home/user/teg-plus/docs/compras/SOLICITACAO_COMPRAS_PARCIAL.docx")
print("Salvo parcial. Continuando...")
