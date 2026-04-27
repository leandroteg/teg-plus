"""Parte 2 — Etapas detalhadas + máquina de estados."""
import sys
sys.path.insert(0, "/home/user/teg-plus/scripts")
from gerar_doc_solicitacao_compras import (  # noqa
    doc, add_heading, add_paragraph, add_bullet, add_table,
    add_code_block, add_callout, page_break, shade_cell, set_cell_borders,
)
from docx.shared import Pt, RGBColor


# 5. ETAPA 1 — CRIAÇÃO
add_heading(doc, "5. Etapa 1 — Criação da Solicitação", level=1)
add_paragraph(doc, "Tela responsável: NovaRequisicao.tsx (1.369 linhas)", italic=True)
add_paragraph(doc, "Hook principal: useCriarRequisicao()", italic=True)
add_paragraph(doc, "Webhook destino: POST /compras/requisicao (n8n)", italic=True)

add_heading(doc, "5.1 Pré-requisitos", level=2)
for item in [
    "Usuário autenticado com role mínima de Requisitante.",
    "Obra ativa cadastrada (relacionada ao centro de custo).",
    "Categoria de material/serviço definida (1 das 12 categorias).",
]:
    add_bullet(doc, item)

add_heading(doc, "5.2 Passos da tela", level=2)
add_table(
    doc,
    ["Passo", "Tela", "Ações", "Validações"],
    [
        ["1", "Categoria",
         "Buscar categoria por nome/código/keywords. IA Parse Helper opcional para extrair dados de PDF/foto.",
         "Categoria obrigatória"],
        ["2", "Detalhes",
         "Solicitante (auto), Obra (lookup), Descrição (UPPERCASE), Justificativa, Urgência, Data necessidade, Centro de Custo (auto pela obra), Classe Financeira",
         "Obra obrigatória; se urgente, justificar"],
        ["3", "Itens",
         "Adicionar linhas (descrição, qtd, unidade, valor estimado, destino). Cola CSV suportada. Autocomplete por estoque.",
         "Mínimo 1 item; qtd > 0"],
        ["4", "Confirmação",
         "Revisar valor total, alçada calculada, mínimo de cotações exigido. Salvar.",
         "Valor total > 0"],
    ],
)

add_heading(doc, "5.3 Cálculos automáticos", level=2)
for item in [
    "valor_estimado = SUM(qty × valor_unitário) por item.",
    "alcada_nivel = determinar_alcada(valor_estimado) → função SQL.",
    "comprador_id = lookup pela categoria (cmp_categorias.comprador_nome).",
    "centro_custo_id = lookup pela obra.",
    "Mínimo de cotações: regra JSONB em cmp_categorias.cotacoes_regras.",
]:
    add_bullet(doc, item)

add_heading(doc, "5.4 IA Parse (opcional)", level=2)
add_paragraph(
    doc,
    "O solicitante pode arrastar um PDF, imagem (foto de uma cotação informal, lista do encarregado, "
    "anotação) ou colar texto. O sistema chama o webhook /compras/requisicao-ai (Gemini Flash Vision), "
    "que retorna:",
)
add_code_block(
    doc,
    """{
  "itens": [{ "descricao": "...", "quantidade": 500, "unidade": "m", "valor_unitario_estimado": 45.50 }],
  "obra_sugerida": "SE Frutal",
  "urgencia_sugerida": "normal",
  "categoria_sugerida": "MATERIAIS_OBRA",
  "comprador_sugerido": { "id": "comp-1", "nome": "Lauany" },
  "confianca": 0.92
}""",
)

add_heading(doc, "5.5 Saída — RC criada", level=2)
add_paragraph(
    doc,
    "Ao confirmar, o n8n gera número (formato RC-YYYYMM-NNNNN), insere em cmp_requisicoes "
    "+ cmp_requisicao_itens, cria primeiro registro em apr_aprovacoes (Alçada 1) com token único "
    "e dispara notificação WhatsApp/e-mail ao aprovador.",
)
add_callout(
    doc,
    "Status inicial",
    "rascunho (local) → pendente (após salvar) → em_aprovacao (após n8n criar a primeira aprovação)",
)

page_break(doc)

# 6. ETAPA 2 — APROVAÇÃO
add_heading(doc, "6. Etapa 2 — Aprovação Multi-Alçada", level=1)
add_paragraph(doc, "Tela responsável: AprovAi.tsx (1.856 linhas)", italic=True)
add_paragraph(doc, "Hook principal: useAprovacoesPendentes() / useDecisaoRequisicao()", italic=True)
add_paragraph(doc, "Webhook destino: POST /compras/aprovacao (n8n)", italic=True)

add_heading(doc, "6.1 Estrutura de Alçadas", level=2)
add_table(
    doc,
    ["Nível", "Cargo", "Faixa de Valor", "Prazo SLA", "Aprovador Padrão"],
    [
        ["1", "Coordenador", "Até R$ 5.000", "24h", "Welton / Claudinor"],
        ["2", "Gerente", "R$ 5.001 a R$ 25.000", "48h", "Laucídio"],
        ["3", "Diretor", "R$ 25.001 a R$ 100.000", "72h", "Diretor Executivo"],
        ["4", "CEO", "Acima de R$ 100.000", "72h", "CEO"],
    ],
)
add_paragraph(
    doc,
    "Função SQL: determinar_alcada(valor DECIMAL) — retorna 1 a 4 conforme faixa.",
    italic=True,
)

add_heading(doc, "6.2 Fluxo Sequencial", level=2)
add_paragraph(
    doc,
    "Uma RC com valor que exige Alçada 3 não pula direto para Alçada 3 — passa por todas as "
    "alçadas intermediárias (1 → 2 → 3). Cada aprovador toma decisão isoladamente, e o n8n só "
    "cria o próximo registro em apr_aprovacoes após a decisão da alçada anterior.",
)

add_heading(doc, "6.3 Decisões Possíveis", level=2)
add_table(
    doc,
    ["Decisão", "Efeito", "Status RC", "Próxima Etapa"],
    [
        ["✅ Aprovada", "Atualiza apr_aprovacoes; cria próxima alçada (se houver)", "em_aprovacao ou aprovada", "Próxima alçada ou cotação"],
        ["❌ Rejeitada", "Encerra fluxo; notifica solicitante", "rejeitada", "Terminal"],
        ["❓ Esclarecimento", "Preserva alçada; notifica solicitante para responder", "em_esclarecimento", "Solicitante responde, volta à mesma alçada"],
    ],
)

add_heading(doc, "6.4 Aprovação por Token (AprovAi Mobile)", level=2)
add_paragraph(
    doc,
    "Para reduzir fricção, cada aprovação gera um token único (apr_aprovacoes.token). O aprovador "
    "recebe via WhatsApp um link no formato:",
)
add_code_block(doc, "https://teg-plus.app/aprovacao?token=apr-f5e3a1d8-4b2e-...")
add_paragraph(
    doc,
    "Ao acessar, o aprovador vê o card da requisição sem precisar fazer login. O token é de uso "
    "único e expira após a decisão ou após o SLA. Útil para diretores/CEO em viagem.",
)

add_callout(
    doc,
    "Alerta de cotação inadequada",
    "Antes de aprovar, o sistema chama get_alerta_cotacao(requisicao_id) e exibe banner se "
    "o número de cotações coletadas for inferior à política da categoria. O aprovador pode "
    "aprovar mesmo assim (com justificativa) ou pedir esclarecimento.",
)

page_break(doc)

# 7. ETAPA 3 — COTAÇÃO
add_heading(doc, "7. Etapa 3 — Cotação (RFQ)", level=1)
add_paragraph(doc, "Tela responsável: CotacaoForm.tsx (1.862 linhas)", italic=True)
add_paragraph(doc, "Hook principal: useCotacao() / useFinalizarCotacao()", italic=True)

add_heading(doc, "7.1 Quando inicia", level=2)
add_paragraph(
    doc,
    "Após a última alçada aprovar, o n8n cria automaticamente um registro em cmp_cotacoes "
    "com status 'pendente', vinculado ao comprador da categoria. A RC fica com status 'em_cotacao'. "
    "O comprador recebe notificação e acessa FilaCotacoes.tsx → CotacaoForm.tsx.",
)

add_heading(doc, "7.2 Política de Mínimo de Fornecedores", level=2)
add_table(
    doc,
    ["Faixa de Valor", "Mínimo de Cotações", "Fonte"],
    [
        ["Até R$ 500", "1 fornecedor", "cmp_categorias.cotacoes_regras"],
        ["R$ 501 a R$ 2.000", "2 fornecedores", "cmp_categorias.cotacoes_regras"],
        ["Acima de R$ 2.000", "3 fornecedores", "cmp_categorias.cotacoes_regras"],
    ],
)
add_paragraph(
    doc,
    "Regra JSONB armazenada por categoria — permite políticas customizadas (ex.: SERVICOS pode "
    "exigir 3 cotações mesmo abaixo de R$ 2.000).",
    italic=True,
)

add_heading(doc, "7.3 Coleta de Propostas", level=2)
add_paragraph(doc, "Para cada fornecedor, o comprador preenche:")
for item in [
    "Dados do fornecedor (nome, contato, CNPJ, e-mail, telefone) — autocomplete por CNPJ via BrasilAPI.",
    "Valor total da proposta.",
    "Prazo de entrega (em dias).",
    "Condição de pagamento (à vista, 30/60/90 dias, parcelado).",
    "Itens com preço unitário (quantidade vem da RC, é READ-ONLY).",
    "Observações.",
    "Upload do PDF/imagem da cotação (Gemini Vision pode preencher campos automaticamente).",
]:
    add_bullet(doc, item)

add_heading(doc, "7.4 Comparativo de Fornecedores", level=2)
add_paragraph(
    doc,
    "O componente CotacaoComparativo.tsx renderiza tabela lado-a-lado e gráfico de preços. "
    "O comprador marca o fornecedor vencedor (selecionado=true). Suporta SPLIT — itens diferentes "
    "podem ir para fornecedores diferentes (ex.: cabo com Siemens, conectores com Nexans).",
)

add_heading(doc, "7.5 Finalização", level=2)
add_paragraph(doc, "Ao clicar 'Finalizar cotação', o sistema:")
for item in [
    "INSERT em cmp_cotacao_fornecedores (uma linha por fornecedor + itens_precos JSONB).",
    "UPDATE em cmp_cotacoes (status='concluida', fornecedor_selecionado_id, valor_selecionado).",
    "UPDATE em cmp_requisicoes (status='cotacao_aprovada').",
    "Cria nova aprovação tipo='cotacao' (para validação financeira da proposta vencedora).",
    "Notifica o gerente financeiro / próximo aprovador.",
]:
    add_bullet(doc, item)

add_heading(doc, "7.6 Devolução ao Solicitante", level=2)
add_paragraph(
    doc,
    "Se o comprador identificar erro de escopo (item errado, quantidade inviável, descrição "
    "incompleta), pode devolver a RC sem coletar cotações. Status passa a 'devolvida_solicitante', "
    "salva mensagem em devolucao_msg/devolucao_por/devolucao_em. Solicitante revisa e reenvia "
    "via useReenviarAposDevolucao().",
)

page_break(doc)

# 8. ETAPA 4 — PEDIDO
add_heading(doc, "8. Etapa 4 — Emissão do Pedido de Compra (PO)", level=1)
add_paragraph(doc, "Componente: EmitirPedidoModal.tsx | Tela: Pedidos.tsx (1.806 linhas)", italic=True)
add_paragraph(doc, "Hook principal: useEmitirPedido() / usePedidos()", italic=True)

add_heading(doc, "8.1 Geração", level=2)
add_paragraph(
    doc,
    "Após cotação aprovada, o comprador (ou o gerente, dependendo da alçada) clica 'Emitir PO'. "
    "O modal vem pré-preenchido com dados da cotação vencedora.",
)

add_heading(doc, "8.2 Campos do Pedido", level=2)
add_table(
    doc,
    ["Campo", "Origem", "Editável"],
    [
        ["numero_pedido", "Auto (PO-YYYYMM-NNNNN)", "Não"],
        ["fornecedor", "Cotação vencedora", "Não"],
        ["valor_total", "Cotação vencedora", "Não"],
        ["data_pedido", "now()", "Não"],
        ["data_prevista_entrega", "now() + prazo_entrega_dias", "Sim"],
        ["condicao_pagamento", "Cotação vencedora", "Sim"],
        ["parcelas_preview", "Calculado da condição", "Sim (via JSONB)"],
        ["centro_custo / classe_financeira", "Herda da RC", "Sim"],
        ["observacoes", "Vazio", "Sim"],
    ],
)

add_heading(doc, "8.3 Efeitos da Emissão", level=2)
for item in [
    "INSERT em cmp_pedidos (status='emitido').",
    "UPDATE em cmp_requisicoes (status='pedido_emitido').",
    "Cria automaticamente Conta a Pagar em fin_contas_pagar (status='aguardando_nf').",
    "Envia PO em PDF para o fornecedor (e-mail) e WhatsApp (se cadastrado).",
    "Notifica solicitante (RC virou PO).",
    "Sincroniza com Omie via workflow n8n (cria CP no Omie ERP).",
]:
    add_bullet(doc, item)

page_break(doc)

# 9. ETAPA 5 — RECEBIMENTO
add_heading(doc, "9. Etapa 5 — Recebimento e Conferência", level=1)
add_paragraph(doc, "Tela: Pedidos.tsx (módulo recebimento)", italic=True)

add_heading(doc, "9.1 Conferência Física", level=2)
add_paragraph(
    doc,
    "Quando o material chega ao almoxarifado da obra, o almoxarife abre o PO no sistema e:",
)
for item in [
    "Confere quantidade física vs PO/NF.",
    "Confere itens (descrição, marca, especificação).",
    "Registra divergências (faltas, sobras, avarias).",
    "Faz upload da NF-e (PDF) e preenche nf_numero.",
    "Clica 'Confirmar Recebimento'.",
]:
    add_bullet(doc, item)

add_heading(doc, "9.2 Efeitos no Sistema", level=2)
for item in [
    "UPDATE cmp_pedidos (status='entregue', data_entrega_real, nf_numero, nf_url).",
    "UPDATE cmp_requisicoes (status='entregue').",
    "INSERT em est_movimentacoes (tipo='entrada_pedido') — alimenta saldo de estoque.",
    "UPDATE fin_contas_pagar (status='liberado' ou 'aguardando_pagamento').",
    "Notifica solicitante (material disponível).",
]:
    add_bullet(doc, item)

add_callout(
    doc,
    "Divergências fiscais",
    "Se quantidade ou valor da NF divergir do PO, o sistema bloqueia o avanço para pagamento "
    "e cria pendência para conciliação manual pela tesouraria.",
)

page_break(doc)

# 10. ETAPA 6 — PAGAMENTO
add_heading(doc, "10. Etapa 6 — Pagamento", level=1)
add_paragraph(doc, "Módulo: Financeiro / Contas a Pagar (fin_contas_pagar)", italic=True)

add_heading(doc, "10.1 Liberação", level=2)
add_paragraph(
    doc,
    "A tesouraria acessa /financeiro/contas-pagar, valida NF + recebimento, e clica "
    "'Liberar Pagamento'. O status do CP passa a 'liberado' e do PO a 'liberado'.",
)

add_heading(doc, "10.2 Registro do Pagamento", level=2)
add_paragraph(doc, "Ao efetivar o pagamento (PIX, TED, boleto):")
for item in [
    "Tesouraria registra forma, valor, data e comprovante.",
    "UPDATE fin_contas_pagar (status='pago', data_pagamento, comprovante_url).",
    "UPDATE cmp_pedidos (status_pagamento='pago', pago_em).",
    "UPDATE cmp_requisicoes (status='pago') — STATUS FINAL.",
    "Sincroniza baixa com Omie ERP via n8n.",
    "Notifica solicitante e gestor da obra (ciclo concluído).",
]:
    add_bullet(doc, item)

page_break(doc)

# 11. MÁQUINA DE ESTADOS
add_heading(doc, "11. Máquina de Estados (Status)", level=1)
add_paragraph(
    doc,
    "Todos os status possíveis de uma RC, quem promove a transição e qual a próxima ação esperada.",
)

add_table(
    doc,
    ["Status", "Quem Define", "Ação Seguinte", "Cor (UI)"],
    [
        ["rascunho", "Solicitante (local)", "Salvar para 'pendente'", "Cinza"],
        ["pendente", "Sistema (após salvar)", "Aguarda primeira alçada", "Laranja"],
        ["em_aprovacao", "Sistema (roteado)", "Aguarda alçada atual", "Azul"],
        ["em_esclarecimento", "Aprovador", "Solicitante responde", "Âmbar"],
        ["devolvida_solicitante", "Comprador", "Solicitante revisa + reenvia", "Vermelho claro"],
        ["aprovada", "Sistema (última alçada)", "Encaminhada para cotação", "Verde claro"],
        ["rejeitada", "Aprovador", "Terminal", "Vermelho"],
        ["em_cotacao", "Sistema", "Comprador cota", "Azul"],
        ["cotacao_enviada", "Comprador", "Aguarda análise/aprovação", "Azul claro"],
        ["cotacao_aprovada", "Sistema/Gerente", "Emitir PO", "Verde"],
        ["cotacao_rejeitada", "Gerente Financeiro", "Refazer cotação", "Vermelho claro"],
        ["pedido_emitido", "Sistema (PO criada)", "Aguarda entrega", "Roxo"],
        ["em_entrega", "Sistema (NF recebida)", "Almoxarife confere", "Roxo claro"],
        ["entregue", "Almoxarife", "Aguarda pagamento", "Verde claro"],
        ["aguardando_pgto", "Sistema", "Tesouraria libera/paga", "Amarelo"],
        ["pago", "Tesouraria", "Terminal (sucesso)", "Verde escuro"],
        ["cancelada", "Admin/Comprador", "Terminal", "Cinza"],
    ],
)

doc.save("/home/user/teg-plus/docs/compras/SOLICITACAO_COMPRAS_PARCIAL.docx")
print("Parte 2/3 concluída.")
