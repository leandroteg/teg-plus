"""Tutorial do usuário — Como fazer uma Solicitação de Compras no TEG+."""
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- Helpers ----------
def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_heading(doc, text, level=1, color=(0x1F, 0x4E, 0x79)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_callout(doc, title, body, color="DEEBF6", title_color=(0x1F, 0x4E, 0x79)):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, color)
    set_cell_borders(cell)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(*title_color)
    r2 = p.add_run(body)
    r2.font.size = Pt(10)
    doc.add_paragraph("")  # espaço


def add_step_box(doc, number, title, content_lines):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(1.5)
    table.columns[1].width = Cm(15)
    num_cell = table.rows[0].cells[0]
    shade_cell(num_cell, "1F4E79")
    set_cell_borders(num_cell)
    num_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    num_cell.text = ""
    p = num_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(number))
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    body_cell = table.rows[0].cells[1]
    shade_cell(body_cell, "F4F8FB")
    set_cell_borders(body_cell)
    body_cell.text = ""
    title_p = body_cell.paragraphs[0]
    tr = title_p.add_run(title)
    tr.bold = True
    tr.font.size = Pt(14)
    tr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    for line in content_lines:
        p = body_cell.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)
    doc.add_paragraph("")


def add_table(doc, headers, rows, header_color="1F4E79", zebra="F2F2F2"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        shade_cell(hdr[i], header_color)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                shade_cell(cells[c_idx], zebra)
            set_cell_borders(cells[c_idx])
    return table


def page_break(doc):
    doc.add_page_break()


# ---------- Build ----------
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


# ============ CAPA ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("\n\n\nGuia Rápido")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t2.add_run("Como fazer uma\nSolicitação de Compras")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t3.add_run("\nTEG+ ERP")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph("\n\n\n\n\n")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Material de apoio para o solicitante\n")
r.italic = True
r.font.size = Pt(13)
r2 = meta.add_run(f"Versão 1.0  •  {datetime.now().strftime('%d/%m/%Y')}")
r2.italic = True
r2.font.size = Pt(11)
r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

page_break(doc)


# ============ INTRODUÇÃO ============
add_heading(doc, "Antes de começar", level=1)
add_paragraph(
    doc,
    "Este guia explica, passo a passo, como abrir uma Solicitação de Compras (RC) no TEG+. "
    "Use sempre que precisar comprar materiais, EPIs, ferramentas, serviços ou contratar locações "
    "para a sua obra ou setor.",
)
add_paragraph(doc, "")
add_paragraph(doc, "O que você vai precisar ter em mãos:", bold=True)
add_bullet(doc, "Saber qual obra ou centro de custo será debitado.")
add_bullet(doc, "Lista clara dos itens (descrição, quantidade, unidade).")
add_bullet(doc, "Uma estimativa do valor (mesmo que aproximada).")
add_bullet(doc, "Data em que você precisa do material/serviço.")
add_bullet(doc, "Justificativa, se a compra for urgente.")

add_callout(
    doc,
    "Dica",
    "Se você tem uma cotação informal, foto de uma lista do encarregado ou um PDF do "
    "fornecedor, o TEG+ consegue ler o documento e preencher quase tudo automaticamente "
    "usando IA. Veja o atalho no Passo 2.",
    color="FFF4D6",
    title_color=(0xB8, 0x86, 0x0B),
)

page_break(doc)


# ============ FLUXO RESUMIDO ============
add_heading(doc, "Visão geral do fluxo", level=1)
add_paragraph(
    doc,
    "Sua solicitação passa pelas etapas abaixo. Você participa principalmente da Etapa 1 "
    "(criação) e, se necessário, da Etapa 2 (responder esclarecimentos).",
)
add_paragraph(doc, "")

add_table(
    doc,
    ["Etapa", "Quem faz", "O que acontece"],
    [
        ["1. Criação", "Você (solicitante)", "Preenche a RC e envia"],
        ["2. Aprovação", "Seu coordenador / gerente", "Aprova ou pede esclarecimento"],
        ["3. Cotação", "Comprador (Lauany / Fernando / Aline)", "Coleta preços de fornecedores"],
        ["4. Pedido", "Comprador", "Emite o Pedido de Compra"],
        ["5. Entrega", "Almoxarifado", "Recebe e confere o material"],
        ["6. Pagamento", "Tesouraria", "Paga o fornecedor"],
    ],
)

add_callout(
    doc,
    "Você será notificado",
    "A cada mudança de etapa importante, você receberá uma notificação por WhatsApp e/ou "
    "e-mail (ex.: quando for aprovada, quando o pedido for emitido, quando o material chegar).",
)

page_break(doc)


# ============ PASSO 1 ============
add_heading(doc, "Passo a passo", level=1)
add_paragraph(
    doc,
    "Acesse o TEG+ e clique no menu lateral em \"Compras\" → \"Nova Requisição\".",
    italic=True,
)
add_paragraph(doc, "")

add_step_box(
    doc, 1, "Escolha a categoria",
    [
        "Selecione o tipo de compra que você precisa fazer. As categorias mais usadas são:",
        "•  Materiais de Obra — cabos, conectores, postes, ferragens, etc.",
        "•  EPI / EPC — capacetes, luvas, óculos, cintos, fitas de sinalização.",
        "•  Ferramental — alicates, chaves, marteletes, escadas.",
        "•  Frota e Equipamentos — peças, manutenção, combustível.",
        "•  Serviços — terceirizados, consultorias, manutenções.",
        "•  Locação — equipamentos, veículos, andaimes.",
        "•  Mobilização / Alojamento / Alimentação — apoio à equipe em campo.",
        "•  Escritório — papelaria, impressão, suprimentos.",
        "",
        "👉 Use o campo de busca no topo se não encontrar a categoria desejada.",
    ],
)

add_callout(
    doc,
    "Por que isso importa?",
    "A categoria define quem vai cuidar da sua compra (qual comprador), quantas cotações "
    "serão exigidas e qual aprovador vai recebê-la. Escolha com atenção.",
)


# ============ PASSO 2 ============
add_step_box(
    doc, 2, "Preencha os detalhes",
    [
        "Os principais campos são:",
        "•  Solicitante: já vem preenchido com seu nome.",
        "•  Obra: selecione a obra/canteiro que será debitado.",
        "•  Descrição: explique resumidamente o que está sendo solicitado.",
        "•  Justificativa: explique POR QUE você precisa (ajuda na aprovação).",
        "•  Urgência: Normal / Urgente / Crítica.",
        "•  Data necessidade: quando você precisa receber o material/serviço.",
        "•  Centro de Custo: já vem preenchido automaticamente pela obra.",
        "",
        "⚡ Atalho com IA: se você tem um documento (PDF, foto da lista, e-mail do "
        "fornecedor), arraste-o no campo \"IA Parse\" e o sistema vai extrair os itens, "
        "quantidades e valores automaticamente. Revise antes de avançar.",
    ],
)

add_callout(
    doc,
    "Atenção à urgência",
    "Se marcar \"Urgente\" ou \"Crítica\", você é OBRIGADO a justificar por que. "
    "Compras urgentes sem justificativa adequada podem ser devolvidas pelo aprovador.",
    color="FDE7E7",
    title_color=(0xC0, 0x39, 0x2B),
)

page_break(doc)


# ============ PASSO 3 ============
add_step_box(
    doc, 3, "Adicione os itens",
    [
        "Para cada item que você precisa, clique em \"+ Adicionar item\" e preencha:",
        "•  Descrição: nome do item (ex.: CABO XLPE 50MM²).",
        "•  Quantidade: quanto você precisa.",
        "•  Unidade: un, kg, m, m², m³, L, pç, etc.",
        "•  Valor unitário estimado: preço aproximado (pode ser estimativa).",
        "•  Destino: Estoque, Patrimônio ou Nenhum (uso direto).",
        "",
        "💡 Dica: você pode COLAR uma tabela do Excel direto na lista de itens — "
        "o TEG+ entende o formato automaticamente.",
        "",
        "💡 Se o item já existe no almoxarifado, comece a digitar e o sistema vai "
        "sugerir o item cadastrado (autocomplete).",
    ],
)


# ============ PASSO 4 ============
add_step_box(
    doc, 4, "Revise e envie",
    [
        "Antes de confirmar, o sistema mostra um resumo:",
        "•  Valor total calculado.",
        "•  Alçada de aprovação que será necessária.",
        "•  Quantas cotações serão exigidas pelo comprador.",
        "•  Quem é o comprador responsável.",
        "",
        "Confira tudo. Se estiver correto, clique em \"Confirmar e Enviar\".",
        "",
        "Pronto! Sua RC ganha um número (formato RC-AAAAMM-NNNNN) e segue para aprovação.",
    ],
)

add_callout(
    doc,
    "Você poderá acompanhar tudo",
    "Vá em \"Compras\" → \"Minhas Requisições\" para ver o status de cada RC, em qual "
    "etapa está e quem é o responsável atual.",
)

page_break(doc)


# ============ APÓS O ENVIO ============
add_heading(doc, "Depois que você envia", level=1)

add_heading(doc, "✅ Aprovada", level=2)
add_paragraph(
    doc,
    "Sua RC vai automaticamente para o comprador da categoria. Ele coletará as "
    "cotações e seguirá com a emissão do pedido.",
)

add_heading(doc, "❓ Esclarecimento solicitado", level=2)
add_paragraph(
    doc,
    "O aprovador pode pedir mais informações. Você receberá uma notificação. "
    "Acesse a RC, leia a mensagem do aprovador no bloco amarelo e digite sua "
    "resposta. Após enviar, a RC volta para o mesmo aprovador.",
)

add_heading(doc, "↩️ Devolvida pelo comprador", level=2)
add_paragraph(
    doc,
    "Se o comprador identificar algum erro de escopo (item errado, quantidade "
    "inviável, descrição incompleta), ele pode devolver para você ajustar. "
    "Edite a RC e reenvie.",
)

add_heading(doc, "❌ Rejeitada", level=2)
add_paragraph(
    doc,
    "Se a RC for rejeitada, ela é encerrada. Para tentar novamente, crie uma nova "
    "RC ajustando os pontos apontados pelo aprovador na justificativa da rejeição.",
)

page_break(doc)


# ============ PERGUNTAS FREQUENTES ============
add_heading(doc, "Perguntas frequentes", level=1)

add_heading(doc, "Posso editar uma RC depois de enviada?", level=2)
add_paragraph(
    doc,
    "Apenas se ela voltar para você — quando o aprovador pede esclarecimento ou o "
    "comprador devolve. Em outros casos, entre em contato com o admin de Compras "
    "para cancelar e criar uma nova.",
)

add_heading(doc, "E se eu não souber o valor exato?", level=2)
add_paragraph(
    doc,
    "Sem problema. Coloque uma estimativa razoável. O comprador vai coletar os "
    "preços reais e o valor final pode ser diferente. Mas evite chutes muito "
    "abaixo do real, pois isso muda a alçada de aprovação.",
)

add_heading(doc, "Quanto tempo demora para aprovar?", level=2)
add_paragraph(
    doc,
    "Depende do valor:",
)
add_bullet(doc, "Até R$ 5.000 — até 24 horas (1 alçada).")
add_bullet(doc, "R$ 5.001 a R$ 25.000 — até 48 horas (2 alçadas).")
add_bullet(doc, "R$ 25.001 a R$ 100.000 — até 72 horas (3 alçadas).")
add_bullet(doc, "Acima de R$ 100.000 — até 72 horas (4 alçadas, inclui CEO).")

add_heading(doc, "Posso enviar uma RC com vários fornecedores diferentes?", level=2)
add_paragraph(
    doc,
    "Sim. Você cria UMA solicitação com todos os itens. O comprador é quem decide "
    "se compra tudo de um único fornecedor ou divide entre vários (split).",
)

add_heading(doc, "Posso anexar fotos ou desenhos?", level=2)
add_paragraph(
    doc,
    "Sim. Use o campo de IA Parse no Passo 2 para anexar PDFs/fotos. Eles ficam "
    "registrados na RC e ajudam o aprovador e o comprador a entenderem o pedido.",
)

add_heading(doc, "E se for muito urgente, mesmo?", level=2)
add_paragraph(
    doc,
    "Marque \"Crítica\" e justifique muito bem. Avise também o comprador por "
    "WhatsApp/telefone. O sistema notifica o aprovador imediatamente, mas "
    "uma comunicação direta acelera ainda mais.",
)

add_heading(doc, "Por que minha RC foi devolvida?", level=2)
add_paragraph(
    doc,
    "Os motivos mais comuns são: descrição vaga, quantidade incorreta, item "
    "duplicado, obra errada, falta de justificativa, valor muito fora da realidade. "
    "Leia a mensagem do comprador, ajuste e reenvie.",
)

page_break(doc)


# ============ DICAS ============
add_heading(doc, "Boas práticas", level=1)
add_paragraph(doc, "Para acelerar a aprovação e a entrega:", bold=True)
add_paragraph(doc, "")
add_bullet(doc, "Seja específico na descrição. Ex.: \"CABO XLPE 50MM² 0,6/1KV PRETO\" é melhor que \"cabo elétrico\".")
add_bullet(doc, "Informe a marca quando a especificação técnica for crítica.")
add_bullet(doc, "Inclua a unidade correta (m vs un vs pç fazem diferença na hora de cotar).")
add_bullet(doc, "Estime valores realistas — facilita a alçada e a vida do comprador.")
add_bullet(doc, "Justifique a urgência se for o caso (\"parada de obra\", \"risco de multa\", \"segurança da equipe\").")
add_bullet(doc, "Concentre itens na mesma RC quando vão para a mesma obra/categoria.")
add_bullet(doc, "Não duplique RCs — se já enviou, acompanhe pelo \"Minhas Requisições\".")
add_bullet(doc, "Responda rapidamente quando aparecer pedido de esclarecimento — isso destrava o fluxo.")

page_break(doc)


# ============ AJUDA ============
add_heading(doc, "Precisa de ajuda?", level=1)
add_paragraph(doc, "Em caso de dúvidas:")
add_paragraph(doc, "")
add_bullet(doc, "Sobre a categoria, comprador ou aprovação: fale com seu coordenador direto.")
add_bullet(doc, "Sobre erro no sistema, login ou acesso: contate o suporte de TI.")
add_bullet(doc, "Sobre status da entrega: consulte em \"Pedidos\" no menu lateral.")
add_bullet(doc, "Sobre pagamento ao fornecedor: financeiro / tesouraria.")

add_paragraph(doc, "")
add_paragraph(
    doc,
    "Bom uso do TEG+!",
    bold=True, size=14, color=(0x1F, 0x4E, 0x79),
)


# Salva
final_path = "/home/user/teg-plus/docs/compras/TUTORIAL_USUARIO_SOLICITACAO_COMPRAS.docx"
doc.save(final_path)
print(f"Tutorial gerado: {final_path}")
